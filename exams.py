from tkinter import *
from tkinter import filedialog
import docx


top = Tk()



fpath = ''

def getfilename(e):
    filename=filedialog.askopenfilename(filetypes=(("MS WORD","*.docx"), ("TEXT","*.txt")))
    e.insert(0,filename)
    global fpath
    fpath = filename
    print(fpath)

def save(optn, *data):
    print(fpath[-4 :])
    if (fpath[-4 :] == 'docx') :
        mydoc = docx.Document(fpath)
        if(optn == 1):
            print("mcq")
            qtn = mydoc.add_paragraph('Q. ' + data[0])
            qtn.add_run('a) ' + data[1])
            qtn.add_run('b) ' + data[2])
            qtn.add_run('c) ' + data[3])
            qtn.add_run('d) ' + data[4])
            mydoc.save(fpath)
        elif(optn == 2):
            print("nq")          
            qtn = mydoc.add_paragraph('Q. ' + data[0])
            mydoc.save(fpath)
    else :
        f = open(fpath, 'a')
        if(optn == 1):
            print("mcq")
            f.write('\n\nQ. ' + data[0])
            f.write('a) ' + data[1])
            f.write('b) ' + data[2])
            f.write('c) ' + data[3])
            f.write('d) ' + data[4])
            f.close()
        elif(optn == 2):
            print("nq")          
            f.write('\n\nQ. ' + data[0])
            f.close()
        






def mcqAdd():
    mcqlevel = Toplevel()
    mcqText = Label ( mcqlevel, text = 'Enter Question')
    mcqText.grid(row = 0, columnspan = 2)
    mcqQuestion = Text(mcqlevel, bd =5,  height = 5, width = 70)
    mcqQuestion.grid(row = 1, columnspan = 2)
    mcqText1 = Label ( mcqlevel, text = 'Enter Options')
    mcqText1.grid(row = 2, columnspan = 2)
    optnText1 = Label ( mcqlevel, text = 'a)')
    optnText1.grid(row = 3, column = 0)
    optn1 = Text(mcqlevel, bd =5,  height = 3, width = 35)
    optn1.grid(row = 3, column = 1)
    optnText2 = Label ( mcqlevel, text = 'b)')
    optnText2.grid(row = 4, column = 0)
    optn2 = Text(mcqlevel, bd =5,  height = 3, width = 35)
    optn2.grid(row = 4, column = 1)
    optnText3 = Label ( mcqlevel, text = 'c}')
    optnText3.grid(row = 5, column = 0)
    optn3 = Text(mcqlevel, bd =5,  height = 3, width = 35)
    optn3.grid(row = 5, column = 1)
    optnText4 = Label ( mcqlevel, text = 'd)')
    optnText4.grid(row = 6, column = 0)
    optn4 = Text(mcqlevel, bd =5,  height = 3, width = 35)
    optn4.grid(row = 6, column = 1)
    savebtn = Button(mcqlevel, text ="SAVE", command=lambda:save(1, mcqQuestion.get(1.0,END),optn1.get(1.0,END), optn2.get(1.0,END), optn3.get(1.0,END), optn4.get(1.0,END)))
    savebtn.grid(row = 7, columnspan = 2)
    

def nqAdd():
    nqlevel = Toplevel()
    nqText = Label ( nqlevel, text = 'Enter Question')
    nqText.grid(row = 0, column = 0)
    nqQuestion = Text(nqlevel, bd =5,  height = 5, width = 70)
    nqQuestion.grid(row = 1, column = 0)
    savebtn = Button(nqlevel, text ="SAVE", command=lambda:save(2, nqQuestion.get(1.0,END)))
    savebtn.grid(row = 2, column = 0)
    



p=Label(top,text="File").grid(row=0)
e=Entry(top)
e.grid(row=0,column=1)
fb=Button(top,text="GETFILE",command=lambda:getfilename(e))
fb.grid(row=0,column=2)
mcqbtn = Button(top, text ="MCQ", command = mcqAdd)
nqbtn = Button(top, text ="Noraml Question", command = nqAdd)
mcqbtn.grid(row=1,column=0)
nqbtn.grid(row=1, column=1)
top.mainloop()
